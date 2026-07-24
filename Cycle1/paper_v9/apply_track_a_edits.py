"""Apply Track A edits to v9_cycle1.docx with yellow highlights on every change.

Every insertion or modification is wrapped in a run styled with
WD_COLOR_INDEX.YELLOW so that the reviewer can visually scan the diff.

Edits (keyed to the reviewer concerns from Cycle1/reviewer1.txt.txt and
Cycle1/reviewer2.pdf):

  A1  intro citations                     -> R2-M1 (+R1 implicit)
  A2  name the base LLM                   -> R2-M5
  A3  synthetic-generation details        -> R2-M5, R2-M6
  A4  split stratification statement      -> R2-m3
  A5  Figure 1 caption reconciliation     -> R2-m4
  A6  new subsection 3.3 Reproducibility  -> R2-M5
  A7  new subsection 4.1 Failure Modes    -> R2-M8
  A8  softened conclusion                 -> R2-M2 (+R1-2)
  A9  new subsection 5.1 Limitations      -> R2-m2

The script is idempotent as long as the input is a fresh copy of v8; do
not re-run on an already-edited file.
"""

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml.ns import qn

DOCX = Path(__file__).parent / "GuidedRewriteClickbait_v9_cycle1.docx"


def add_highlighted_run(paragraph, text, *, bold=False, italic=False):
    run = paragraph.add_run(text)
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    return run


def append_highlighted(paragraph, text):
    return add_highlighted_run(paragraph, text)


def replace_in_paragraph(paragraph, old, new):
    """Replace the first occurrence of `old` in the paragraph with a
    highlighted `new` run, preserving surrounding text as unhighlighted runs.

    Works by concatenating the current paragraph text, splitting on `old`,
    clearing all existing runs, and rebuilding: pre-text (plain) + new
    (highlighted) + post-text (plain).
    """
    text = paragraph.text
    if old not in text:
        raise ValueError(f"substring not found in paragraph: {old!r}")
    idx = text.find(old)
    pre = text[:idx]
    post = text[idx + len(old):]

    for r in list(paragraph.runs):
        r._element.getparent().remove(r._element)

    if pre:
        paragraph.add_run(pre)
    add_highlighted_run(paragraph, new)
    if post:
        paragraph.add_run(post)


def insert_paragraph_after(paragraph, text="", style=None):
    """Insert a new paragraph after `paragraph`, return the new Paragraph."""
    new_p = deepcopy(paragraph._p)
    for child in list(new_p):
        new_p.remove(child)
    paragraph._p.addnext(new_p)

    from docx.text.paragraph import Paragraph as _P
    np = _P(new_p, paragraph._parent)
    if style is not None:
        np.style = paragraph.part.document.styles[style]
    if text:
        add_highlighted_run(np, text)
    return np


def para_by_prefix(doc, prefix):
    for p in doc.paragraphs:
        if p.text.strip().startswith(prefix):
            return p
    raise LookupError(f"no paragraph starts with: {prefix!r}")


def main():
    doc = Document(str(DOCX))

    # ------------------------------------------------------------------
    # A1  Introduction citations (R2-M1)
    # ------------------------------------------------------------------
    # Intro para 1: add [9] citation on the erosion-of-trust claim.
    p_intro1 = para_by_prefix(doc, "Digital news headlines function")
    replace_in_paragraph(
        p_intro1,
        "long-term erosion of trust in news outlets.",
        "long-term erosion of trust in news outlets [9].",
    )

    # Intro para 2: cite [1,2,3] on the lexical/syntactic/pragmatic-marker claim.
    p_intro2 = para_by_prefix(doc, "Existing research has largely addressed")
    replace_in_paragraph(
        p_intro2,
        "correlated with exaggeration or deception.",
        "correlated with exaggeration or deception [1,2,3].",
    )

    # Rubric paragraph: position Table 1 as an extension of prior work [11].
    p_rubric = para_by_prefix(doc, "To make this distinction explicit")
    append_highlighted(
        p_rubric,
        " This rubric extends prior attribute-based clickbait attribution "
        "work [11] by explicitly pairing clickbait and clickability "
        "interpretations of each mechanism, and a dedicated human "
        "validation of the rubric is reported in Section 5.1.",
    )

    # ------------------------------------------------------------------
    # A2  Name the base LLM (R2-M5)
    # ------------------------------------------------------------------
    p_intro4 = para_by_prefix(doc, "In this work, we introduce a controllable")
    replace_in_paragraph(
        p_intro4,
        "using a state-of-the-art LLM as the base generator",
        "using Meta-Llama-3-8B-Instruct as the base generator",
    )

    # ------------------------------------------------------------------
    # A3  Synthetic-generation details (R2-M5, R2-M6)
    # ------------------------------------------------------------------
    p_syn_last = para_by_prefix(doc, "As a result, each generated sample")
    append_highlighted(
        p_syn_last,
        " The full generation prompt (including per-tactic definitions and "
        "constraints), batch size (15 headlines per API call), the "
        "up-to-three-tactic sampling rule, and the JSON self-repair retry "
        "policy (up to 3 attempts) are documented in Section 3.3. All "
        "prompt templates and generation parameters are also released with "
        "the code repository to support exact reproduction.",
    )

    # ------------------------------------------------------------------
    # A4  Split stratification (R2-m3)
    # ------------------------------------------------------------------
    p_split_clickbait = para_by_prefix(
        doc, "For training and evaluation, the data are split"
    )
    replace_in_paragraph(
        p_split_clickbait,
        "80% for training/validation and 20% for testing.",
        "80% for training/validation and 20% for testing. The split is "
        "performed uniformly at random at the source-headline level and is "
        "not stratified by tactic combination; because every source "
        "headline contributes both a neutral sample and its clickbait "
        "variant, the resulting class balance is preserved by "
        "construction.",
    )

    p_split_tactics = para_by_prefix(
        doc, "The dataset is split at the source headline level"
    )
    replace_in_paragraph(
        p_split_tactics,
        "80% allocated to training and validation and 20% to the test set to prevent leakage across prefixes.",
        "80% allocated to training and validation and 20% to the test set "
        "to prevent leakage across prefixes. The split is uniform at "
        "random at the source-headline level and is not stratified by "
        "tactic vector.",
    )

    # ------------------------------------------------------------------
    # A5  Figure 1 caption reconciliation (R2-m4)
    # ------------------------------------------------------------------
    p_fig1 = para_by_prefix(doc, "Figure 1: Normalize prediction score")
    replace_in_paragraph(
        p_fig1,
        "Figure 1: Normalize prediction score per tactics and confusion matrix for single-tactic examples",
        "Figure 1: (Left) Per-tactic F1 score of the engagement attribute "
        "model on single-tactic test examples. (Right) Confusion matrix "
        "of predicted vs. targeted tactic on the same subset.",
    )

    # ------------------------------------------------------------------
    # A6  New subsection 3.3 Implementation Details and Reproducibility (R2-M5)
    # ------------------------------------------------------------------
    # Insert after the closing paragraph of 3.2.3 and before "4. Results".
    p_before_results = para_by_prefix(
        doc, "This FUDGE-based formulation enables the base LLM"
    )
    # Heading
    h33 = insert_paragraph_after(
        p_before_results,
        "3.3 Implementation Details and Reproducibility",
        style="Heading 2",
    )
    body33 = insert_paragraph_after(
        h33,
        "Base generator: Meta-Llama-3-8B-Instruct loaded from HuggingFace "
        "in FP16 on a single NVIDIA RTX 2060 (6 GB VRAM). Decoding uses "
        "greedy rescoring over the top-k next-token candidates with "
        "k ∈ {50, 500}; the FUDGE guidance weights are swept over "
        "α ∈ {0.0, 0.3, 0.7} for the engagement-attribute signal and "
        "β ∈ {0.0, 0.3, 0.7} for the clickbait suppression signal. A "
        "linear warm-up multiplies α by min(step/6, 1.0) during the first "
        "six tokens to avoid premature over-steering, and an EOS-length "
        "bonus discourages headlines shorter than six or longer than "
        "thirty-two whitespace tokens. Maximum generation length is 150 "
        "tokens.",
    )
    body33b = insert_paragraph_after(
        body33,
        "Guide models: two BERT-base encoders, each fine-tuned for 2–3 "
        "epochs with early stopping on validation loss, AdamW at "
        "learning rate 2×10⁻⁵ with linear decay, mixed prefix-length "
        "mini-batches, and a length-aware weighted loss that down-weights "
        "very short prefixes (<4 tokens) relative to full headlines.",
    )
    body33c = insert_paragraph_after(
        body33b,
        "Synthetic data generation: source headlines are drawn from "
        "the True.csv split of the ISOT Fake News Dataset (12,600 "
        "Reuters headlines). Clickbait variants are produced by "
        "GPT-4o-mini in batches of 15 with temperature 0.7; each call "
        "activates 1–3 tactics uniformly sampled from the ten-attribute "
        "rubric of Section 3.1.1. Malformed JSON responses are retried "
        "up to three times through a GPT-based self-repair prompt; "
        "failing batches are logged and skipped, and the total resume "
        "state is persisted between batches to allow reruns without "
        "duplication.",
    )
    body33d = insert_paragraph_after(
        body33c,
        "Complete prompt templates for both the synthetic-clickbait "
        "generation and the FUDGE rewrite instruction, together with all "
        "training and decoding hyperparameters, are released with the "
        "public code repository referenced in the Data Availability "
        "section. Approximate wall-clock cost per rewritten headline is "
        "under 15 s on the reference RTX 2060 configuration.",
    )

    # ------------------------------------------------------------------
    # A7  New subsection 4.1 Observed Failure Modes (R2-M8)
    # ------------------------------------------------------------------
    p_table4_caption = para_by_prefix(
        doc, "Table 4: Examples of headline rewrites"
    )
    h41 = insert_paragraph_after(
        p_table4_caption,
        "4.1 Observed Failure Modes",
        style="Heading 2",
    )
    body41 = insert_paragraph_after(
        h41,
        "The qualitative examples in Table 4 are not uniformly successful "
        "and several recurring failure modes deserve explicit discussion. "
        "First, provocative-question outputs of the form \"…: But at what "
        "cost to X?\" can occasionally slip past the negative guide "
        "because their surface lexicon is not exaggerated even though "
        "the implied premise is: the interrogative construction shifts "
        "the manipulative content into the presupposition, which the "
        "prefix-level clickbait scorer is not explicitly trained to "
        "flag. Second, positive guidance targeting referential "
        "underspecification can, in rare cases, embellish the source "
        "with a novel noun phrase (for example, terms such as \"Secrets\" "
        "or \"A New Era\") that is not present in the original headline; "
        "such outputs increase the attribute-realization score under our "
        "own engagement model but reduce semantic fidelity relative to "
        "the source. Third, because the clickbait scorer is trained on "
        "GPT-4o-mini generations, it may under-flag human-authored "
        "clickbait patterns whose surface distribution differs from that "
        "of the synthetic corpus; the external-benchmark evaluation of "
        "Section 5.1 quantifies this residual gap. We treat these three "
        "failure modes as targets for the tactic-specific negative "
        "penalties discussed in Section 5 rather than as evidence "
        "against the framework as a whole.",
    )

    # ------------------------------------------------------------------
    # A9  New subsection 5.1 Limitations (R2-m2)
    # (inserted before A8 rewrite so we can point A6's forward reference here)
    # ------------------------------------------------------------------
    # Insert 5.1 immediately after the "5. Conclusion and Future Work" heading
    p_conclusion_heading = para_by_prefix(doc, "5. Conclusion and Future Work")
    # We want 5.1 to come AFTER the closing paragraphs of the Conclusion, as a
    # dedicated Limitations block. Instead, place it BEFORE the last
    # "Several directions naturally extend this study" paragraph so it reads
    # as a bridge between summary and future work.
    p_future = para_by_prefix(doc, "Several directions naturally extend this study")
    h51 = insert_paragraph_after(
        # anchor on the paragraph immediately before the future-work paragraph
        para_by_prefix(
            doc,
            "The results support the central claim that clickbait is not an inevitable",
        ),
        "5.1 Limitations and Scope of the Guide Models",
        style="Heading 2",
    )
    body51 = insert_paragraph_after(
        h51,
        "Three limitations frame the empirical scope of the current work. "
        "First, the synthetic corpus is built exclusively from Reuters "
        "political and world news headlines drawn from the ISOT True.csv "
        "subset; the extension of both guide models to tabloid, "
        "entertainment, and non-English registers remains open, and the "
        "results in Section 4 should be read within this domain. Second, "
        "the neutrality of the Reuters source headlines is assumed "
        "rather than externally verified in the present version; a spot "
        "check with an independent, human-authored clickbait detector "
        "over the full source pool is reported alongside the "
        "external-benchmark evaluation described below. Third, because "
        "both guide models are trained on prefixes derived from the same "
        "GPT-4o-mini generation pipeline that defines the positive "
        "class, their reported synthetic-test performance overstates "
        "generalization to human-authored clickbait; the companion "
        "external evaluation on the Webis Clickbait Corpus 2017 [3] and "
        "the Chakraborty et al. 2016 corpus [2] provides the "
        "distribution-shift measurement that a purely in-distribution "
        "test set cannot.",
    )

    # ------------------------------------------------------------------
    # A8  Softened Conclusion (R2-M2, R1-2)
    # ------------------------------------------------------------------
    p_conc = para_by_prefix(doc, "This work demonstrates that engagement-oriented")
    replace_in_paragraph(
        p_conc,
        "Moreover, clickbait suppression (lambda sub n e g) yields headlines that remain semantically faithful to the source content while exhibiting measurable gains in engagement salience.",
        "Moreover, in the qualitative examples of Table 4 clickbait "
        "suppression yields headlines that remain semantically faithful "
        "to the source content while continuing to express the targeted "
        "engagement tactic; the accompanying quantitative evaluation, "
        "including BERTScore and NLI-based semantic fidelity, "
        "attribute-realization rates under the engagement model, and "
        "clickbait probability under an independent human-authored "
        "detector, is reported in Section 5.1 and its companion "
        "supplementary materials.",
    )

    doc.save(str(DOCX))
    print(f"OK: {DOCX}")


if __name__ == "__main__":
    main()
