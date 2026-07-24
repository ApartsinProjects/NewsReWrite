"""Render v9_cycle1.docx to HTML, preserving yellow highlights as <mark>.

The output is a self-contained HTML file (no external assets) with a
readable single-column layout, a light academic style, and every
yellow-highlighted DOCX run wrapped in a <mark> element so a reader can
scan the Track A diff at a glance.

This is not a general-purpose docx-to-html converter; it handles only
the constructs present in the paper (paragraphs and Heading 1/2/3).
"""

import html
from pathlib import Path

from docx import Document
from docx.enum.text import WD_COLOR_INDEX

DOCX = Path(__file__).parent / "GuidedRewriteClickbait_v9_cycle1.docx"
HTML = Path(__file__).parent / "GuidedRewriteClickbait_v9_cycle1.html"

CSS = """
:root {
  color-scheme: light dark;
  --bg: #ffffff;
  --fg: #101418;
  --muted: #5b6470;
  --accent: #b45309;
  --rule: #e4e7eb;
  --mark: #fff59d;
  --mark-fg: #202020;
}
@media (prefers-color-scheme: dark) {
  :root {
    --bg: #0f1216;
    --fg: #e6e8eb;
    --muted: #9aa3ad;
    --accent: #f0a044;
    --rule: #232830;
    --mark: #b58900;
    --mark-fg: #101010;
  }
}
* { box-sizing: border-box; }
body {
  margin: 0;
  font-family: "Georgia", "Cambria", "Times New Roman", serif;
  color: var(--fg);
  background: var(--bg);
  line-height: 1.55;
  font-size: 17px;
}
main { max-width: 780px; margin: 0 auto; padding: 3rem 1.5rem 6rem; }
h1, h2, h3 {
  font-family: "Helvetica Neue", "Inter", system-ui, sans-serif;
  line-height: 1.25;
  color: var(--fg);
}
h1.title { font-size: 1.9rem; margin: 0 0 .3rem; }
.authors { color: var(--muted); font-family: sans-serif; font-size: .95rem; margin-bottom: .3rem; }
.affil { color: var(--muted); font-family: sans-serif; font-size: .85rem; margin-bottom: .2rem; }
h1.section { font-size: 1.35rem; margin-top: 2.6rem; padding-bottom: .3rem; border-bottom: 1px solid var(--rule); }
h2.subsection { font-size: 1.10rem; margin-top: 2rem; }
h3.subsubsection { font-size: 1.00rem; margin-top: 1.4rem; color: var(--muted); }
p { margin: 0 0 1rem; text-align: justify; hyphens: auto; }
mark {
  background: var(--mark);
  color: var(--mark-fg);
  padding: .05em .15em;
  border-radius: 3px;
}
.callout {
  font-family: sans-serif;
  font-size: .85rem;
  color: var(--muted);
  border-left: 3px solid var(--accent);
  padding: .6rem .9rem;
  margin: 1.4rem 0;
  background: color-mix(in srgb, var(--accent) 6%, transparent);
}
.figref, .tableref { font-style: italic; color: var(--muted); }
"""

CALLOUT_HTML = """
<div class="callout">
  <strong>Cycle 1 revision (v9_cycle1).</strong> Passages highlighted in
  <mark>yellow</mark> mark Track A revisions applied in response to the
  Reviewer 1 and Reviewer 2 reports. A full changelog is provided in
  <code>Cycle1/CHANGELOG_cycle1.md</code>; the point-by-point response
  is in <code>Cycle1/response_letter_cycle1.md</code>.
</div>
"""


def runs_to_html(paragraph):
    parts = []
    for r in paragraph.runs:
        text = html.escape(r.text)
        if not text:
            continue
        if r.font.highlight_color == WD_COLOR_INDEX.YELLOW:
            text = f"<mark>{text}</mark>"
        if r.bold:
            text = f"<strong>{text}</strong>"
        if r.italic:
            text = f"<em>{text}</em>"
        parts.append(text)
    return "".join(parts)


def render_paragraph(paragraph, is_first_body=False):
    text = paragraph.text.strip()
    if not text:
        return ""
    style = paragraph.style.name
    inner = runs_to_html(paragraph)
    if style == "Heading 1":
        cls = "section"
        return f'<h1 class="{cls}">{inner}</h1>'
    if style == "Heading 2":
        return f'<h2 class="subsection">{inner}</h2>'
    if style == "Heading 3":
        return f'<h3 class="subsubsection">{inner}</h3>'
    # Table / figure captions marked by prefix.
    if text.startswith(("Table ", "Figure ")):
        return f'<p class="tableref">{inner}</p>'
    return f"<p>{inner}</p>"


def main():
    doc = Document(str(DOCX))

    title_html = html.escape(doc.paragraphs[1].text.strip())
    authors_html = html.escape(doc.paragraphs[2].text.strip())
    affil1_html = html.escape(doc.paragraphs[4].text.strip())
    affil2_html = html.escape(doc.paragraphs[5].text.strip())
    corr_html = html.escape(doc.paragraphs[7].text.strip())

    body_paras = []
    for i, p in enumerate(doc.paragraphs):
        if i <= 7:
            continue
        rendered = render_paragraph(p)
        if rendered:
            body_paras.append(rendered)

    body = "\n".join(body_paras)

    out = f"""<!doctype html>
<html lang="en">
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>{title_html} — v9 cycle1</title>
<style>{CSS}</style>
</head>
<body>
<main>
<h1 class="title">{title_html}</h1>
<p class="authors">{authors_html}</p>
<p class="affil">{affil1_html}</p>
<p class="affil">{affil2_html}</p>
<p class="affil">{corr_html}</p>
{CALLOUT_HTML}
{body}
</main>
</body>
</html>
"""
    HTML.write_text(out, encoding="utf-8")
    print(f"OK: {HTML}")


if __name__ == "__main__":
    main()
